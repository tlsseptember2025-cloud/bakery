import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
import os
import tempfile
from datetime import datetime
from utils import center_window
from docx.enum.text import WD_ALIGN_PARAGRAPH
from paths import get_db_path


# ----- Optional Word (docx) support -----
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ----- Base paths -----
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RECEIPTS_DIR = os.path.join(BASE_DIR, "receipts")
LOGO_PATH = os.path.join(BASE_DIR, "images", "logo.png")
SIG_PATH = os.path.join(BASE_DIR, "images", "sig.png")


# ===================== HELPERS =====================

def safe_float(value):
    try:
        return float(value)
    except:
        return 0.0


def ensure_receipts_table():
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS receipts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            recipe_name TEXT,
            quantity INTEGER,
            customer_name TEXT,
            total REAL,
            receipt_text TEXT,
            created_at TEXT
        )
    """)
    conn.commit()
    conn.close()


def create_docx_receipt_from_history(
    receipt_id,
    recipe_name,
    qty,
    customer_name,
    total_sale,
    created_at
):
    """
    Create a nicely formatted Word (.docx) receipt using data from history.
    Returns absolute path to the file, or None on failure.
    """
    if not DOCX_AVAILABLE:
        return None

    os.makedirs(RECEIPTS_DIR, exist_ok=True)
    safe_name = "".join(ch for ch in recipe_name if ch.isalnum() or ch in (" ", "_", "-")).strip()

    # Use receipt_id in filename so it's unique and traceable
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    docx_path = os.path.join(RECEIPTS_DIR, f"{safe_name}_#{receipt_id}_{timestamp}.docx")

    doc = Document()

    # Logo (if exists)
    if os.path.exists(LOGO_PATH):
        doc.add_picture(LOGO_PATH, width=Inches(1.2))

    # Title
    title = doc.add_paragraph()
    run = title.add_run("MY BAKERY")
    run.bold = True
    title.alignment = 1  # center

    subtitle = doc.add_paragraph("Fresh Cakes & Pastries")
    subtitle.alignment = 1

    doc.add_paragraph("")  # blank line

    # Basic info
    info_p = doc.add_paragraph()
    info_p.add_run("Receipt #: ").bold = True
    info_p.add_run(str(receipt_id))
    info_p.add_run("\nDate: ").bold = True
    info_p.add_run(created_at)
    info_p.add_run("\nCustomer: ").bold = True
    info_p.add_run(customer_name if customer_name else "N/A")

    doc.add_paragraph("")  # blank line

    # Table with order details
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Qty"
    hdr_cells[2].text = "Price"
    hdr_cells[3].text = "Total"

    row_cells = table.add_row().cells
    row_cells[0].text = recipe_name
    row_cells[1].text = str(qty)
    price_each = (safe_float(total_sale) / qty) if qty else 0.0
    row_cells[2].text = f"{price_each:.2f}"
    row_cells[3].text = f"{safe_float(total_sale):.2f}"

    doc.add_paragraph("")  # blank line

    total_p = doc.add_paragraph()
    run_total = total_p.add_run(f"Grand Total: {safe_float(total_sale):.2f}")
    run_total.bold = True
    doc.add_paragraph("")
    p = doc.add_paragraph("Return/Refund Policy:")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1 = doc.add_paragraph("1. You have to present the original receipt")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph("2. You have 7 days from date of delivery")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3 = doc.add_paragraph("3. Items should be in good new condition")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4 = doc.add_paragraph("4. No refund of any kind!")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    footer = doc.add_paragraph("Thank you for your purchase!")

    if os.path.exists(SIG_PATH):
        doc.add_picture(SIG_PATH, width=Inches(3))

    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.alignment = 1

    doc.save(docx_path)

    if os.path.exists(docx_path):
        return docx_path
    return None


def open_for_printing_from_history(
    receipt_id,
    recipe_name,
    qty,
    customer_name,
    total_sale,
    created_at,
    receipt_text,
    parent
):
    """
    Instead of printing directly, open the receipt in Word or Notepad
    so the user can use the standard print dialog and choose printer.
    """
    if os.name != "nt":
        messagebox.showinfo(
            "Open for Printing",
            "Automatic opening is best supported on Windows.\n"
            "On other systems, please open the file manually and print.",
            parent=parent
        )
        return

    # 1) Try to create & open DOCX in Word (with logo, nice layout)
    docx_path = None
    if DOCX_AVAILABLE:
        try:
            docx_path = create_docx_receipt_from_history(
                receipt_id,
                recipe_name,
                qty,
                customer_name,
                total_sale,
                created_at
            )
        except Exception as e:
            messagebox.showerror("Error", f"Could not create Word receipt:\n{e}", parent=parent)
            docx_path = None

    if docx_path and os.path.exists(docx_path):
        try:
            os.startfile(os.path.abspath(docx_path))  # open Word; user uses File > Print
            return
        except Exception:
            messagebox.showinfo(
                "Open for Printing",
                "Could not open the Word receipt. I will open a text version instead.",
                parent=parent
            )

    # 2) Fallback: open a temporary text file in Notepad
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as tmp:
            tmp.write(receipt_text)
            txt_path = tmp.name
    except Exception as e:
        messagebox.showerror("Error", f"Could not create temporary text file:\n{e}", parent=parent)
        return

    try:
        os.startfile(txt_path)  # opens in Notepad; user uses File > Print
    except Exception as e:
        messagebox.showerror("Error", f"Could not open receipt file:\n{e}", parent=parent)


# ===================== MAIN WINDOW =====================

def open_receipts_history():
    ensure_receipts_table()

    win = tk.Toplevel()
    win.title("Receipts History")
    center_window(win, 900, 600)
    win.lift()
    win.focus_force()

    tk.Label(win, text="Receipts History", font=("Arial", 16, "bold")).pack(pady=10)

    # ---------- FRAME FOR TABLE ----------
    table_frame = tk.Frame(win)
    table_frame.pack(fill="both", expand=True, padx=10, pady=5)

    scroll_y = tk.Scrollbar(table_frame)
    scroll_y.pack(side="right", fill="y")

    scroll_x = tk.Scrollbar(table_frame, orient="horizontal")
    scroll_x.pack(side="bottom", fill="x")

    tree = ttk.Treeview(
        table_frame,
        yscrollcommand=scroll_y.set,
        xscrollcommand=scroll_x.set,
        selectmode="browse"
    )
    tree.pack(fill="both", expand=True)
    scroll_y.config(command=tree.yview)
    scroll_x.config(command=tree.xview)

    tree["columns"] = ("ID", "Recipe Name", "Qty", "Customer", "Total", "Created At")
    tree.column("#0", width=0, stretch=tk.NO)
    tree.column("ID", anchor=tk.CENTER, width=60)
    tree.column("Recipe Name", anchor=tk.W, width=220)
    tree.column("Qty", anchor=tk.CENTER, width=60)
    tree.column("Customer", anchor=tk.W, width=180)
    tree.column("Total", anchor=tk.CENTER, width=90)
    tree.column("Created At", anchor=tk.CENTER, width=160)

    tree.heading("#0", text="")
    tree.heading("ID", text="ID")
    tree.heading("Recipe Name", text="Recipe Name")
    tree.heading("Qty", text="Qty")
    tree.heading("Customer", text="Customer")
    tree.heading("Total", text="Total")
    tree.heading("Created At", text="Created At")

    tree.tag_configure('evenrow', background='#f0f0ff')
    tree.tag_configure('oddrow', background='#ffffff')

    # ---------- LOAD RECEIPTS ----------
    def load_receipts():
        for row in tree.get_children():
            tree.delete(row)

        conn = sqlite3.connect(get_db_path())
        cur = conn.cursor()
        try:
            cur.execute("""
                SELECT id, recipe_name, quantity, customer_name, total, created_at
                FROM receipts
                ORDER BY datetime(created_at) DESC
            """)
            rows = cur.fetchall()
        except sqlite3.OperationalError:
            conn.close()
            messagebox.showerror(
                "Error",
                "The 'receipts' table does not exist.\nMake a cake first or ensure DB setup.",
                parent=win
            )
            return

        conn.close()

        for idx, (rid, recipe_name, qty, customer, total, created_at) in enumerate(rows):
            tag = 'evenrow' if idx % 2 == 0 else 'oddrow'
            customer_display = customer if customer else "N/A"
            tree.insert(
                "",
                tk.END,
                values=(rid, recipe_name, qty, customer_display, f"{safe_float(total):.2f}", created_at),
                tags=(tag,)
            )

    load_receipts()

    # ---------- PREVIEW FUNCTION ----------
    def preview_selected_receipt():
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select a receipt to view.", parent=win)
            win.lift()
            win.focus_force()
            return

        item = tree.item(selected[0])
        receipt_id = item['values'][0]

        conn = sqlite3.connect(get_db_path())
        cur = conn.cursor()
        cur.execute("""
            SELECT recipe_name, receipt_text, created_at, quantity, customer_name, total
            FROM receipts
            WHERE id = ?
        """, (receipt_id,))
        row = cur.fetchone()
        conn.close()

        if not row:
            messagebox.showerror("Error", "Receipt not found in database.", parent=win)
            win.lift()
            win.focus_force()
            return

        recipe_name, receipt_text, created_at, qty, customer, total = row

        # ----- PREVIEW WINDOW -----
        preview = tk.Toplevel(win)
        preview.title(f"Receipt #{receipt_id} - {recipe_name}")
        center_window(preview, 900, 900)

        preview.transient(win)
        preview.grab_set()
        preview.lift()
        preview.attributes("-topmost", True)
        preview.after(150, lambda: preview.attributes("-topmost", False))
        preview.focus_force()

        tk.Label(
            preview,
            text=f"Receipt #{receipt_id}",
            font=("Arial", 14, "bold")
        ).pack(pady=5)

        top_info = f"Recipe: {recipe_name}    |    Qty: {qty}    |    Total: {safe_float(total):.2f}"
        tk.Label(
            preview,
            text=top_info,
            font=("Arial", 10)
        ).pack(pady=2)

        tk.Label(
            preview,
            text=f"Customer: {customer if customer else 'N/A'}",
            font=("Arial", 10)
        ).pack(pady=2)

        tk.Label(
            preview,
            text=f"Date: {created_at}",
            font=("Arial", 10)
        ).pack(pady=2)

        text_box = tk.Text(
            preview,
            font=("Courier New", 11),
            wrap="word",
            bg="#f9f9f9"
        )
        text_box.pack(fill="both", expand=True, padx=10, pady=10)
        text_box.insert("1.0", receipt_text)
        text_box.config(state="disabled")

        # Buttons: Open for Printing + Close
        btn_frame = tk.Frame(preview)
        btn_frame.pack(pady=8)

        tk.Button(
            btn_frame,
            text="Open for Printing",
            width=18,
            command=lambda: open_for_printing_from_history(
                receipt_id,
                recipe_name,
                qty,
                customer,
                total,
                created_at,
                receipt_text,
                preview
            )
        ).grid(row=0, column=0, padx=5)

        tk.Button(
            btn_frame,
            text="Close",
            width=12,
            command=preview.destroy
        ).grid(row=0, column=1, padx=5)

        preview.wait_window()

    # ---------- DOUBLE-CLICK TO PREVIEW ----------
    def on_double_click(event):
        preview_selected_receipt()

    tree.bind("<Double-1>", on_double_click)

    # ---------- BUTTONS ----------
    btn_frame = tk.Frame(win)
    btn_frame.pack(pady=10)

    tk.Button(
        btn_frame,
        text="View Selected Receipt",
        width=22,
        command=preview_selected_receipt
    ).grid(row=0, column=0, padx=5)

    tk.Button(
        btn_frame,
        text="Refresh List",
        width=15,
        command=load_receipts
    ).grid(row=0, column=1, padx=5)

    tk.Button(
        btn_frame,
        text="Close",
        width=12,
        command=win.destroy
    ).grid(row=0, column=2, padx=5)


