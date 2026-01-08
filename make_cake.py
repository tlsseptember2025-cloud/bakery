import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
import os
import tempfile
from datetime import datetime
from utils import center_window
from paths import get_db_path


# ----- Optional Word (docx) support -----
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ----- Base paths -----
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RECEIPTS_DIR = os.path.join(BASE_DIR, "receipts")
LOGO_PATH = os.path.join(BASE_DIR, "images", "logo.png")


# ===================== HELPERS =====================

def safe_float(value):
    try:
        return float(value)
    except:
        return 0.0


def ensure_receipts_table():
    """Make sure receipts table exists with needed columns."""
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS receipts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            recipe_name TEXT,
            quantity INTEGER,
            customer_name TEXT,
            total REAL,
            receipt_text TEXT,
            created_at TEXT
        )
    """)
    conn.commit()
    conn.close()


def save_receipt_to_db(recipe_name, qty, customer_name, total, receipt_text, created_at):
    ensure_receipts_table()
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO receipts (recipe_name, quantity, customer_name, total, receipt_text, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (
        recipe_name,
        qty,
        customer_name,
        total,
        receipt_text,
        created_at
    ))
    conn.commit()
    conn.close()


def ensure_cost_column():
    """Ensure ingredients has cost_per_unit column (for possible internal reports)."""
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("PRAGMA table_info(ingredients)")
    cols = [r[1].lower() for r in cur.fetchall()]
    if "cost_per_unit" not in cols:
        cur.execute("ALTER TABLE ingredients ADD COLUMN cost_per_unit REAL DEFAULT 0")
        conn.commit()
    conn.close()


# ===================== RECEIPT GENERATION =====================

def generate_text_receipt(recipe_name, qty, customer_name, selling_price, ingredients_used):
    """
    Creates the TEXT receipt, saves it, and returns:
      file_path, receipt_text, total_sale, created_at
    """
    os.makedirs(RECEIPTS_DIR, exist_ok=True)

    # Safe filename (no weird chars)
    safe_name = "".join(ch for ch in recipe_name if ch.isalnum() or ch in (" ", "_", "-")).strip()
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    file_path = os.path.join(RECEIPTS_DIR, f"{safe_name}_{timestamp}.txt")

    total_sale = selling_price * qty
    created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Layout: wider, with top padding and left margin
    page_width = 80          # characters
    left_margin = " " * 6    # spaces from left edge
    top_padding_lines = 4    # blank lines at top for some space

    def center(text: str) -> str:
        return left_margin + text.center(page_width)

    lines = []

    # Top blank space
    for _ in range(top_padding_lines):
        lines.append("")

    lines.append(left_margin + "=" * page_width)
    lines.append(center("MY BAKERY"))
    lines.append(center("Fresh Cakes & Pastries"))
    lines.append(left_margin + "=" * page_width)

    lines.append(left_margin + f"Date     : {created_at}")
    lines.append(left_margin + f"Customer : {customer_name if customer_name else 'N/A'}")
    lines.append(left_margin + "-" * page_width)

    lines.append(left_margin + f"Item     : {recipe_name}")
    lines.append(left_margin + f"Quantity : {qty}")
    lines.append(left_margin + f"Price    : {selling_price:.2f}")
    lines.append(left_margin + f"TOTAL    : {total_sale:.2f}")

    lines.append(left_margin + "-" * page_width)
    lines.append(left_margin + "Ingredients Used:")
    for name, qty_used, unit, cpu in ingredients_used:
        qty_used = safe_float(qty_used)
        lines.append(left_margin + f"- {name}: {qty_used} {unit}")

    lines.append(left_margin + "-" * page_width)
    lines.append(center("Conditions for ReturnsTo be eligible for a return"))
    lines.append(center("Conditions for ReturnsTo be eligible for a return"))
    lines.append(center("Conditions for ReturnsTo be eligible for a return"))
    lines.append(center("Conditions for ReturnsTo be eligible for a return"))
    lines.append(center("Conditions for ReturnsTo be eligible for a return"))
    lines.append(left_margin + "=" * page_width)
    lines.append(left_margin + "-" * page_width)
    lines.append(center("Thank you! Please visit again."))
    lines.append(left_margin + "=" * page_width)

    receipt_text = "\n".join(lines)

    # Save text file
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(receipt_text)

    # Save to DB
    save_receipt_to_db(recipe_name, qty, customer_name, total_sale, receipt_text, created_at)

    return file_path, receipt_text, total_sale, created_at


def create_docx_receipt(recipe_name, qty, customer_name, total_sale, created_at):
    """
    Create a nicely formatted Word (.docx) receipt with logo and return its absolute path.
    Requires python-docx to be installed.
    """
    if not DOCX_AVAILABLE:
        return None

    os.makedirs(RECEIPTS_DIR, exist_ok=True)
    safe_name = "".join(ch for ch in recipe_name if ch.isalnum() or ch in (" ", "_", "-")).strip()
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    docx_path = os.path.join(RECEIPTS_DIR, f"{safe_name}_{timestamp}.docx")

    doc = Document()

    # Logo (if exists)
    if os.path.exists(LOGO_PATH):
        doc.add_picture(LOGO_PATH, width=Inches(1.2))

    # Title
    title = doc.add_paragraph()
    run = title.add_run("MY BAKERY")
    run.bold = True
    title.alignment = 1  # center

    subtitle = doc.add_paragraph("Fresh Cakes & Pastries")
    subtitle.alignment = 1

    doc.add_paragraph("")  # blank line

    # Basic info
    info_p = doc.add_paragraph()
    info_p.add_run("Date: ").bold = True
    info_p.add_run(created_at)
    info_p.add_run("\nCustomer: ").bold = True
    info_p.add_run(customer_name if customer_name else "N/A")

    doc.add_paragraph("")  # blank line

    # Table with order details
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Qty"
    hdr_cells[2].text = "Price"
    hdr_cells[3].text = "Total"

    row_cells = table.add_row().cells
    row_cells[0].text = recipe_name
    row_cells[1].text = str(qty)
    row_cells[2].text = f"{total_sale / qty:.2f}" if qty else "0.00"
    row_cells[3].text = f"{total_sale:.2f}"

    doc.add_paragraph("")  # blank line

    total_p = doc.add_paragraph()
    run_total = total_p.add_run(f"Grand Total: {total_sale:.2f}")
    run_total.bold = True

    doc.add_paragraph("")  # blank line
    doc.add_paragraph("Conditions for ReturnsTo be eligible for a return")
    doc.add_paragraph("Conditions for ReturnsTo be eligible for a return")
    doc.add_paragraph("Conditions for ReturnsTo be eligible for a return")
    doc.add_paragraph("Conditions for ReturnsTo be eligible for a return")
    doc.add_paragraph("Conditions for ReturnsTo be eligible for a return")
    doc.add_paragraph("")
    doc.add_paragraph("")
    footer = doc.add_paragraph("Thank you for your purchase!")
    footer.alignment = 1

    doc.save(docx_path)

    if os.path.exists(docx_path):
        return docx_path
    return None


# ===================== OPEN FOR PRINTING & PREVIEW =====================

def open_for_printing(file_txt_path, receipt_text, parent,
                      recipe_name, qty, customer_name, total_sale, created_at):
    """
    Instead of printing silently, open the receipt in Word or Notepad
    so the user can use the normal Print dialog and choose the printer.
    """
    if os.name != "nt":
        messagebox.showinfo(
            "Open for Printing",
            "Automatic opening is best supported on Windows.\n"
            "On other systems, please open the file manually and print.",
            parent=parent
        )
        return

    # 1) Try to create & open DOCX in Word (nicer, with logo)
    docx_path = None
    if DOCX_AVAILABLE:
        try:
            docx_path = create_docx_receipt(recipe_name, qty, customer_name, total_sale, created_at)
        except Exception as e:
            messagebox.showerror("Error", f"Could not create Word receipt:\n{e}", parent=parent)
            docx_path = None

    if docx_path and os.path.exists(docx_path):
        try:
            os.startfile(os.path.abspath(docx_path))  # just open (user prints from Word)
            return
        except Exception as e:
            messagebox.showinfo(
                "Open for Printing",
                "Could not open Word receipt. I will open the text receipt instead.",
                parent=parent
            )

    # 2) Fallback: open text receipt in Notepad
    path_to_open = os.path.abspath(file_txt_path) if file_txt_path else None
    if not path_to_open or not os.path.exists(path_to_open):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as tmp:
                tmp.write(receipt_text)
                path_to_open = tmp.name
        except Exception as e:
            messagebox.showerror("Error", f"Could not create temporary text file:\n{e}", parent=parent)
            return

    try:
        os.startfile(path_to_open)  # opens in Notepad → user uses File > Print
    except Exception as e:
        messagebox.showerror("Error", f"Could not open receipt file:\n{e}", parent=parent)


def show_receipt_preview(parent, receipt_text, file_path,
                         recipe_name, qty, customer_name, total_sale, created_at):
    """Preview window with 'Open for Printing' button."""
    preview = tk.Toplevel(parent)
    preview.title("Receipt Preview")
    center_window(preview, 900, 900)

    preview.transient(parent)
    preview.grab_set()
    preview.lift()
    preview.attributes("-topmost", True)
    preview.after(150, lambda: preview.attributes("-topmost", False))
    preview.focus_force()

    tk.Label(preview, text="Receipt Preview", font=("Arial", 14, "bold")).pack(pady=10)

    text_box = tk.Text(preview, font=("Courier New", 11), wrap="word", bg="#f9f9f9")
    text_box.pack(fill="both", expand=True, padx=10, pady=10)
    text_box.insert("1.0", receipt_text)
    text_box.config(state="disabled")

    btn_frame = tk.Frame(preview)
    btn_frame.pack(pady=10)

    tk.Button(
        btn_frame,
        text="Open for Printing",
        width=18,
        command=lambda: open_for_printing(
            file_path,
            receipt_text,
            preview,
            recipe_name,
            qty,
            customer_name,
            total_sale,
            created_at
        )
    ).grid(row=0, column=0, padx=5)

    tk.Button(
        btn_frame,
        text="Close",
        width=12,
        command=preview.destroy
    ).grid(row=0, column=1, padx=5)

    preview.wait_window()


# ===================== MAIN WINDOW =====================

def open_make_cake():
    ensure_cost_column()
    ensure_receipts_table()

    win = tk.Toplevel()
    win.title("Make Cake")
    center_window(win, 650, 520)
    win.lift()
    win.focus_force()

    tk.Label(win, text="Make Cake", font=("Arial", 16, "bold")).pack(pady=10)

    # --- Recipe selection ---
    tk.Label(win, text="Select Recipe", font=("Arial", 12)).pack()
    recipe_combo = ttk.Combobox(win, width=35, state="readonly")
    recipe_combo.pack(pady=5)

    # --- Quantity ---
    tk.Label(win, text="Quantity (number of cakes)", font=("Arial", 12)).pack()
    qty_entry = tk.Entry(win, width=20)
    qty_entry.pack(pady=5)

    # --- Customer Name ---
    tk.Label(win, text="Customer Name (optional)", font=("Arial", 12)).pack()
    customer_entry = tk.Entry(win, width=35)
    customer_entry.pack(pady=5)

    # Load recipes into combo
    conn = sqlite3.connect(get_db_path())
    cur = conn.cursor()
    cur.execute("SELECT id, name FROM recipes ORDER BY name ASC")
    recipes = cur.fetchall()
    conn.close()

    recipe_map = {name: rid for rid, name in recipes}
    recipe_combo["values"] = list(recipe_map.keys())

    def make_cake():
        recipe_name = recipe_combo.get().strip()
        qty_text = qty_entry.get().strip()
        customer_name = customer_entry.get().strip()

        if not recipe_name or not qty_text:
            messagebox.showerror("Error", "Please select a recipe and enter quantity.", parent=win)
            win.lift(); win.focus_force()
            return

        try:
            qty = int(qty_text)
            if qty <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror("Error", "Quantity must be a positive whole number.", parent=win)
            win.lift(); win.focus_force()
            return

        recipe_id = recipe_map[recipe_name]

        conn = sqlite3.connect(get_db_path())
        cur = conn.cursor()

        # Get selling price
        cur.execute("SELECT selling_price FROM recipes WHERE id = ?", (recipe_id,))
        row = cur.fetchone()
        selling_price = safe_float(row[0]) if row else 0.0

        # Get ingredients needed and inventory + cost
        cur.execute("""
            SELECT 
                ri.ingredient_id,
                ri.quantity,       -- required per 1 cake
                i.quantity,        -- available in stock
                i.name,
                i.unit,
                i.cost_per_unit
            FROM recipe_ingredients ri
            JOIN ingredients i ON ri.ingredient_id = i.id
            WHERE ri.recipe_id = ?
        """, (recipe_id,))
        rows = cur.fetchall()

        if not rows:
            conn.close()
            messagebox.showerror("Error", "This recipe has no ingredients linked.", parent=win)
            win.lift(); win.focus_force()
            return

        # --- Check stock ---
        for ing_id, req_per_cake, avail_qty, name, unit, cpu in rows:
            req_per_cake = safe_float(req_per_cake)
            avail_qty = safe_float(avail_qty)
            total_required = req_per_cake * qty

            if avail_qty < total_required:
                conn.close()
                messagebox.showerror(
                    "Low Stock",
                    f"Not enough {name}\n\n"
                    f"Available: {avail_qty} {unit}\n"
                    f"Required for {qty} cake(s): {total_required} {unit}",
                    parent=win
                )
                win.lift(); win.focus_force()
                return

        # --- Deduct inventory ---
        used_ingredients_for_receipt = []
        for ing_id, req_per_cake, avail_qty, name, unit, cpu in rows:
            req_per_cake = safe_float(req_per_cake)
            avail_qty = safe_float(avail_qty)
            total_required = req_per_cake * qty
            new_qty = avail_qty - total_required
            if new_qty < 0:
                new_qty = 0

            cur.execute(
                "UPDATE ingredients SET quantity = ? WHERE id = ?",
                (new_qty, ing_id)
            )

            used_ingredients_for_receipt.append((name, total_required, unit, cpu))

        conn.commit()
        conn.close()

        # --- Generate and show receipt (text + openable docx) ---
        file_path, receipt_text, total_sale, created_at = generate_text_receipt(
            recipe_name,
            qty,
            customer_name,
            selling_price,
            used_ingredients_for_receipt
        )

        show_receipt_preview(
            win,
            receipt_text,
            file_path,
            recipe_name,
            qty,
            customer_name,
            total_sale,
            created_at
        )

        # Reset quantity & customer
        qty_entry.delete(0, tk.END)
        customer_entry.delete(0, tk.END)
        win.lift()
        win.focus_force()

    tk.Button(
        win,
        text="Make Cake",
        font=("Arial", 12, "bold"),
        width=22,
        command=make_cake
    ).pack(pady=20)
